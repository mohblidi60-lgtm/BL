from flask import Flask, render_template, request, send_file, jsonify
import json, os, re, sqlite3
from generate import generate_bl

app = Flask(__name__)

# ── Database ───────────────────────────────────────────────
DB_PATH = os.environ.get('DB_PATH', os.path.join(os.path.dirname(__file__), 'data.db'))

def get_db():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    return con

def init_db():
    with get_db() as con:
        con.execute('''CREATE TABLE IF NOT EXISTS clients
                       (code TEXT PRIMARY KEY, nom TEXT, nom1 TEXT, nom2 TEXT)''')

init_db()

def load_clients():
    with get_db() as con:
        rows = con.execute('SELECT code, nom, nom1, nom2 FROM clients ORDER BY nom').fetchall()
        return [dict(r) for r in rows]

def save_clients(clients):
    with get_db() as con:
        con.execute('DELETE FROM clients')
        con.executemany(
            'INSERT OR REPLACE INTO clients (code, nom, nom1, nom2) VALUES (?,?,?,?)',
            [(c.get('code',''), c.get('nom',''), c.get('nom1',''), c.get('nom2','')) for c in clients]
        )

# ── Routes ─────────────────────────────────────────────────
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/clients', methods=['GET'])
def get_clients():
    return jsonify(load_clients())

@app.route('/clients', methods=['POST'])
def post_clients():
    clients = request.get_json()
    save_clients(clients)
    return jsonify({'ok': True, 'count': len(clients)})

@app.route('/generate', methods=['POST'])
def generate():
    data = request.get_json()
    try:
        buf = generate_bl(data)
        bl_num = data.get('bl_num', '???')
        m = re.match(r'(\d{2})(\d{2})-(.*)', bl_num)
        short = f'{m.group(2)}-{m.group(3)}' if m else bl_num
        fname = f'Bon de livraison interne\u5185\u90E8\u53D1\u8D27\u5355 {short}.docx'
        return send_file(
            buf,
            as_attachment=True,
            download_name=fname,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/import', methods=['POST'])
def import_docx():
    from docx import Document as Doc
    import re as re2

    f = request.files.get('file')
    if not f:
        return jsonify({'error': 'Aucun fichier'}), 400

    try:
        doc = Doc(f)

        # ── Header: BL number ──
        bl_num = ''
        for section in doc.sections:
            hdr = section.header
            for para in hdr.paragraphs:
                if 'livraison' in para.text.lower():
                    m = re2.search(r'N[o°º]?\.?\s*([\w\-]+)', para.text)
                    if m:
                        bl_num = m.group(1).strip()
                    break

        # ── Table 0: Infos ──
        t0 = doc.tables[0]
        def cell_txt(row, col):
            try: return t0.rows[row].cells[col].text.strip()
            except: return ''

        date_raw    = cell_txt(0, 1)
        vendeur     = cell_txt(0, 3)
        client      = cell_txt(1, 1)
        code_client = cell_txt(1, 3)

        bl_date = ''
        m = re2.match(r'(\d{1,2})/(\d{1,2})/(\d{4})', date_raw)
        if m:
            bl_date = f'{m.group(3)}-{m.group(2).zfill(2)}-{m.group(1).zfill(2)}'

        # ── Table 1: Articles ──
        t1 = doc.tables[1]

        def clean(s):
            return (s or '').strip()

        def parse_num(s):
            s = clean(s).replace(' ', '').replace('\u202f', '').replace(',', '.')
            try: return float(s)
            except: return 0

        articles = []
        data_rows = t1.rows[2:-3]
        for row in data_rows:
            cells = row.cells
            if len(cells) < 7:
                continue
            desig   = clean(cells[1].text)
            ref     = clean(cells[2].text)
            ml      = clean(cells[3].text)
            rouleau = clean(cells[4].text)
            carton  = clean(cells[5].text)
            prix    = parse_num(cells[6].text)

            if not desig and not ref:
                continue

            # Calculate rpc from rouleau/carton
            rpc_val = 6
            try:
                r = float(rouleau) if rouleau else 0
                c = float(carton)  if carton  else 0
                if c > 0 and r > 0:
                    calc = round(r / c)
                    if calc in [3, 4, 6, 8]:
                        rpc_val = calc
            except:
                pass

            articles.append({
                'desig':   desig,
                'ref':     ref,
                'ml':      ml,
                'rouleau': rouleau,
                'carton':  carton,
                'prix':    str(prix) if prix else '',
                'rpc':     str(rpc_val)
            })

        # ── Totals: last 3 rows (TOTAL / Remise / Net) ──
        total_montant  = 0
        remise_montant = 0
        remise_pct     = 0
        WNS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc'
        try:
            tot_tcs = t1.rows[-3]._tr.findall(WNS)
            rem_tcs = t1.rows[-2]._tr.findall(WNS)
            # text of physical tc
            def tc_text(tc):
                return ''.join(p.text for p in tc.iter()
                               if p.tag.endswith('}t'))
            total_montant  = parse_num(tc_text(tot_tcs[-2]))
            remise_montant = parse_num(tc_text(rem_tcs[-2]))
            if total_montant > 0 and remise_montant > 0:
                remise_pct = round(remise_montant / total_montant * 100, 2)
        except:
            pass

        return jsonify({
            'bl_num':        bl_num,
            'bl_date':       bl_date,
            'vendeur':       vendeur,
            'client':        client,
            'code_client':   code_client,
            'articles':      articles,
            'remise_pct':    remise_pct,
            'remise_montant': remise_montant,
            'total_montant': total_montant
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
