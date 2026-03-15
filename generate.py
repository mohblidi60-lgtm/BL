"""
Génère un bon de livraison en copiant le template original.
"""
import copy
import io
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

TEMPLATE_PATH = "template.docx"

def set_cell_text(cell, text):
    for para in cell.paragraphs:
        rpr = None
        for run in para.runs:
            rpr_el = run._r.find(qn('w:rPr'))
            if rpr_el is not None:
                rpr = copy.deepcopy(rpr_el)
            break
        for r in para._p.findall(qn('w:r')): para._p.remove(r)
        for bm in para._p.findall(qn('w:bookmarkStart')): para._p.remove(bm)
        for bm in para._p.findall(qn('w:bookmarkEnd')): para._p.remove(bm)
        for pe in para._p.findall(qn('w:proofErr')): para._p.remove(pe)
        r_el = etree.SubElement(para._p, qn('w:r'))
        if rpr is not None:
            r_el.insert(0, rpr)
        t_el = etree.SubElement(r_el, qn('w:t'))
        t_el.text = str(text)
        if str(text).startswith(' ') or str(text).endswith(' '):
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        break


def set_tc_text(tc, text):
    """Set text in a w:tc element directly (works with physical cell XML)."""
    for p in tc.findall(qn('w:p')):
        rpr = None
        for r in p.findall(qn('w:r')):
            rpr_el = r.find(qn('w:rPr'))
            if rpr_el is not None:
                rpr = copy.deepcopy(rpr_el)
            break
        for r in p.findall(qn('w:r')): p.remove(r)
        for x in p.findall(qn('w:proofErr')): p.remove(x)
        for x in p.findall(qn('w:bookmarkStart')): p.remove(x)
        for x in p.findall(qn('w:bookmarkEnd')): p.remove(x)
        r_el = etree.SubElement(p, qn('w:r'))
        if rpr is not None:
            r_el.insert(0, rpr)
        t_el = etree.SubElement(r_el, qn('w:t'))
        t_el.text = str(text)
        if str(text).startswith(' ') or str(text).endswith(' '):
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        break


def clone_row_with_values(orig_row, tc_values):
    """
    Clone a row and set values by PHYSICAL tc index.
    tc_values = [(phys_tc_index, text), ...]
    """
    new_tr = copy.deepcopy(orig_row._tr)
    ns14 = 'http://schemas.microsoft.com/office/word/2010/wordml'
    for el in new_tr.iter():
        el.attrib.pop(f'{{{ns14}}}paraId', None)
        el.attrib.pop(f'{{{ns14}}}textId', None)

    tcs = new_tr.findall(qn('w:tc'))
    for idx, text in tc_values:
        if idx < len(tcs):
            set_tc_text(tcs[idx], text)
    return new_tr


def fmt_num(n):
    n = float(n or 0)
    if n == 0:
        return '0'
    s = f'{n:,.2f}'
    s = s.replace(',', ' ').replace('.', ',')
    return s


def generate_bl(data):
    doc = Document(TEMPLATE_PATH)

    # ── 1. Header ──
    for section in doc.sections:
        hdr = section.header
        for para in hdr.paragraphs:
            if 'Bon de livraison' in para.text:
                bl_num = data.get('bl_num', '????')
                p = para._p
                for r in p.findall(qn('w:r')):
                    p.remove(r)

                def mk_run(text, size=32, superscript=False):
                    r = etree.SubElement(p, qn('w:r'))
                    rpr = etree.SubElement(r, qn('w:rPr'))
                    fonts = etree.SubElement(rpr, qn('w:rFonts'))
                    fonts.set(qn('w:ascii'), 'Arial')
                    fonts.set(qn('w:hAnsi'), 'Arial')
                    fonts.set(qn('w:cs'), 'Arial')
                    pos = etree.SubElement(rpr, qn('w:position'))
                    pos.set(qn('w:val'), '2')
                    sz = etree.SubElement(rpr, qn('w:sz'))
                    sz.set(qn('w:val'), str(size))
                    szCs = etree.SubElement(rpr, qn('w:szCs'))
                    szCs.set(qn('w:val'), str(size))
                    if superscript:
                        va = etree.SubElement(rpr, qn('w:vertAlign'))
                        va.set(qn('w:val'), 'superscript')
                    t = etree.SubElement(r, qn('w:t'))
                    t.text = text
                    if text.startswith(' ') or text.endswith(' '):
                        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

                mk_run('Bon de livraison N', 32)
                mk_run('o', 24, superscript=True)
                mk_run('.' + bl_num, 32)
                break

    # ── 2. Table 0 : Infos ──
    t0 = doc.tables[0]
    set_cell_text(t0.rows[0].cells[1], data.get('date', ''))
    set_cell_text(t0.rows[0].cells[3], data.get('vendeur', ''))
    set_cell_text(t0.rows[1].cells[1], data.get('client', ''))
    set_cell_text(t0.rows[1].cells[3], data.get('code_client', ''))

    # ── 3. Table 1 : Articles ──
    t1 = doc.tables[1]
    template_row = t1.rows[2]  # first data row as template

    # Totals
    total_carton  = sum(float(a.get('carton', 0) or 0) for a in data['articles'])
    total_rouleau = sum(float(a.get('rouleau', 0) or 0) for a in data['articles'])
    total_montant = sum(float(a.get('rouleau', 0) or 0) * float(a.get('prix', 0) or 0)
                       for a in data['articles'])
    remise = float(data.get('remise', 0) or 0)
    net    = total_montant - remise

    # Save originals BEFORE modifying doc
    doc_orig = Document(TEMPLATE_PATH)
    t1_orig  = doc_orig.tables[1]
    tot_row_orig    = t1_orig.rows[5]
    remise_row_orig = t1_orig.rows[6]
    net_row_orig    = t1_orig.rows[7]

    # Remove all rows after header (rows 0 and 1)
    tbl_el = t1._tbl
    all_trs = tbl_el.findall(qn('w:tr'))
    for tr in all_trs[2:]:
        tbl_el.remove(tr)

    # Insert article rows
    for i, art in enumerate(data['articles']):
        rouleau = float(art.get('rouleau', 0) or 0)
        prix    = float(art.get('prix', 0) or 0)
        montant = rouleau * prix
        carton  = art.get('carton', '')
        ml      = art.get('ml', '')
        ref     = art.get('ref', '')
        desig   = art.get('desig', '0.8x22mm')

        new_tr = clone_row_with_values(template_row, [
            (0, str(i + 1)),
            (1, desig if 'Chant' in str(desig) else f'Chant PVC Réf: {desig}'),
            (2, str(ref)),
            (3, str(ml)),
            (4, str(int(rouleau) if rouleau == int(rouleau) else rouleau)),
            (5, str(int(carton) if float(carton or 0) == int(float(carton or 0)) else carton)),
            (6, fmt_num(prix)),
            (7, fmt_num(montant)),
        ])
        tbl_el.append(new_tr)

    # ── Verify physical tc count of totals rows ──
    # Row 5 (TOTAL): physical tcs = [bigspan(0-4), carton, TOTAL_label, montant, spacer]
    # => indices:       0              1             2       3           4
    tot_tr = clone_row_with_values(tot_row_orig, [
        (1, str(int(total_carton))),
        (2, 'TOTAL'),
        (3, fmt_num(total_montant)),
    ])
    tbl_el.append(tot_tr)

    # Row 6 (Remise): [bigspan, remise_label(span2), montant, spacer]
    # => indices:      0        1                     2        3
    rem_tr = clone_row_with_values(remise_row_orig, [
        (1, 'Remise'),
        (2, fmt_num(remise)),
    ])
    tbl_el.append(rem_tr)

    # Row 7 (Net): [bigspan, net_label(span2), montant, spacer]
    net_tr = clone_row_with_values(net_row_orig, [
        (1, 'Net à payer'),
        (2, fmt_num(net)),
    ])
    tbl_el.append(net_tr)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
